from __future__ import annotations

import asyncio
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from app.api import routes_documents, routes_exports
from app.core.config import Settings
from app.core.security import AuthContext

class _UploadFile:
    def __init__(self, path: Path) -> None:
        self.filename = path.name
        self.file = path.open("rb")

    async def seek(self, offset: int) -> None:
        self.file.seek(offset)

    def close(self) -> None:
        self.file.close()


class ApiSampleEndToEndTests(unittest.TestCase):
    def test_synthetic_docx_upload_process_quality_and_export_route(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            sample = _write_synthetic_regulation_docx(tmp_path / "synthetic_regulation.docx")
            settings = Settings(data_dir=tmp_path / "data")
            auth = AuthContext(actor="api-smoke-test", tenant_id="tenant-smoke", auth_mode="api_token")
            upload = _UploadFile(sample)
            try:
                with patch.object(routes_documents, "get_settings", return_value=settings):
                    document = asyncio.run(
                        routes_documents.upload_document(
                            upload,
                            institution_name="Smoke Institution",
                            source_system="LOCAL",
                            source_record_id="sample-board",
                            source_file_id="sample-file",
                            profile_id="default-public-institution",
                            auth_context=auth,
                        )
                    )
                    job = routes_documents.process_document(document["document_id"], None, auth)
                    quality = routes_documents.get_quality(document["document_id"], auth)
            finally:
                upload.close()

            with patch.object(routes_exports, "get_settings", return_value=settings):
                exported = routes_exports.export_document(document["document_id"], "jsonl", auth)

            self.assertEqual("completed", job["status"])
            self.assertTrue(quality["passed"])
            self.assertGreater(quality["chunk_count"], 0)
            self.assertTrue(Path(exported.path).is_file())
            self.assertEqual(f"{document['document_id']}.jsonl", exported.filename)


def _write_synthetic_regulation_docx(path: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise unittest.SkipTest("python-docx is not installed") from exc

    doc = Document()
    doc.add_paragraph("합성 규정")
    doc.add_paragraph("제1조(목적) 이 규정은 합성 테스트 문서의 처리 절차를 정함을 목적으로 한다.")
    doc.add_paragraph("제2조(육아휴직) 직원은 승인 절차에 따라 육아휴직을 신청할 수 있다.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "구분"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "육아휴직"
    table.cell(1, 1).text = "승인 후 사용"
    doc.save(path)
    return path


if __name__ == "__main__":
    unittest.main()
