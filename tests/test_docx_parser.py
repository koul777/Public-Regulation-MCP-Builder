from __future__ import annotations

import importlib.util
import tempfile
import unittest
from pathlib import Path

from app.parsers.base import ParserError
from app.parsers.docx_parser import DocxParser
from app.parsers.factory import get_parser


DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None


class DocxParserTests(unittest.TestCase):
    def test_factory_supports_docx_extension(self) -> None:
        self.assertIsInstance(get_parser(Path("sample.docx")), DocxParser)

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx is not installed")
    def test_preserves_paragraph_table_order(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "ordered.docx"
            doc = Document()
            doc.add_paragraph("제1조 목적")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "구분"
            table.cell(0, 1).text = "내용"
            table.cell(1, 0).text = "가"
            table.cell(1, 1).text = "본문"
            doc.add_paragraph("제2조 적용")
            doc.save(path)

            parsed = DocxParser().parse(path, "doc_ordered")

        blocks = parsed.pages[0].blocks
        self.assertEqual([block.type for block in blocks], ["text", "table", "text"])
        self.assertEqual([block.text for block in blocks], ["제1조 목적", "구분 | 내용\n가 | 본문", "제2조 적용"])
        self.assertEqual(parsed.raw_text, "제1조 목적\n구분 | 내용\n가 | 본문\n제2조 적용")

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx is not installed")
    def test_horizontally_merged_cell_is_not_repeated_per_column(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "merged.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=3)
            header = table.cell(0, 0).merge(table.cell(0, 2))
            header.text = "공통 기준"
            table.cell(1, 0).text = "A"
            table.cell(1, 1).text = "B"
            table.cell(1, 2).text = "C"
            doc.save(path)

            parsed = DocxParser().parse(path, "doc_merged")

        table_block = next(block for block in parsed.pages[0].blocks if block.type == "table")
        self.assertEqual(table_block.text, "공통 기준\nA | B | C")

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx is not installed")
    def test_header_part_is_exposed_as_review_metadata_without_reordering_body(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "header.docx"
            doc = Document()
            doc.add_paragraph("본문")
            doc.sections[0].header.paragraphs[0].text = "머리말"
            doc.sections[0].footer.paragraphs[0].text = "꼬리말"
            doc.save(path)

            parsed = DocxParser().parse(path, "doc_header")

        self.assertEqual([block.text for block in parsed.pages[0].blocks], ["본문"])
        self.assertEqual(parsed.metadata["parser_uncertainty_risk_level"], "medium")
        self.assertIn("docx_unparsed_parts", parsed.metadata["parser_uncertainty_flags"])
        self.assertEqual(parsed.metadata["parser_uncertainty_recommendation"], "review_missing_docx_parts")
        self.assertIn("word/header1.xml", parsed.metadata["docx_unparsed_parts"])
        self.assertIn("word/footer1.xml", parsed.metadata["docx_unparsed_parts"])

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx is not installed")
    def test_invalid_docx_raises_parser_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "invalid.docx"
            path.write_bytes(b"not a docx")

            with self.assertRaisesRegex(ParserError, "Failed to parse DOCX file"):
                DocxParser().parse(path, "doc_invalid_docx")


if __name__ == "__main__":
    unittest.main()
