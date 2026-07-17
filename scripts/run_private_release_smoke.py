from __future__ import annotations

import argparse
import asyncio
import subprocess
from contextlib import ExitStack
from dataclasses import dataclass
import json
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Sequence
from unittest.mock import patch
from urllib.parse import urlsplit
from uuid import uuid4


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

import app.main as app_main
from app.api import routes_documents, routes_exports
from app.core.api_audit import api_audit_path
from app.core.config import Settings, get_settings
from app.core.tenant_access import settings_for_tenant


DEFAULT_EXPORT_FORMATS = ("jsonl", "csv", "markdown", "tables_jsonl", "tables_csv", "quality_json", "quality_md")
DEFAULT_SAMPLE_PATH = Path("data") / "private_release_runtime" / "approved_samples" / "regulation_sample.hwp"
EXPORT_EXTENSIONS = {
    "jsonl": "jsonl",
    "csv": "csv",
    "markdown": "md",
    "tables_jsonl": "tables.jsonl",
    "tables_csv": "tables.csv",
    "quality_json": "quality.json",
    "quality_md": "quality.md",
}


@dataclass(frozen=True)
class AsgiResponse:
    status_code: int
    headers: dict[str, str]
    content: bytes

    def json(self) -> Any:
        return json.loads(self.content.decode("utf-8"))


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def run_smoke(
    *,
    sample_path: Path,
    data_dir: Path,
    export_formats: Sequence[str] = DEFAULT_EXPORT_FORMATS,
    actor: str = "private-release-smoke",
    tenant_id: str = "tenant-smoke",
) -> dict[str, Any]:
    run_started_at = datetime.now(timezone.utc)
    settings = Settings(
        app_env="production",
        data_dir=data_dir,
        api_auth_required=True,
        api_auth_token="smoke-token",
        api_audit_enabled=True,
        tenant_storage_isolation=True,
    )
    settings.data_dir.mkdir(parents=True, exist_ok=True)
    headers = {
        "Authorization": "Bearer smoke-token",
        "X-Actor": actor,
        "X-Tenant-Id": tenant_id,
    }
    exports: list[dict[str, Any]] = []
    multipart_body, multipart_content_type = _multipart_upload_body(
        sample_path,
        {
            "institution_name": "Private Release Smoke Institution",
            "source_system": "LOCAL_SMOKE",
            "source_record_id": "smoke-record",
            "source_file_id": sample_path.name,
            "profile_id": "default-public-institution",
        },
    )

    with ExitStack() as stack:
        stack.enter_context(patch.object(app_main, "get_settings", return_value=settings))
        stack.enter_context(patch.object(routes_documents, "get_settings", return_value=settings))
        stack.enter_context(patch.object(routes_exports, "get_settings", return_value=settings))
        app_main.app.dependency_overrides[get_settings] = lambda: settings
        stack.callback(app_main.app.dependency_overrides.pop, get_settings, None)

        denied_response = asyncio.run(
            _asgi_request(
                "POST",
                "/api/documents",
                body=multipart_body,
                headers={"Content-Type": multipart_content_type},
            )
        )
        missing_tenant_response = asyncio.run(
            _asgi_request(
                "POST",
                "/api/documents",
                body=multipart_body,
                headers={
                    "Authorization": "Bearer smoke-token",
                    "X-Actor": actor,
                    "Content-Type": multipart_content_type,
                },
            )
        )
        upload_response = asyncio.run(
            _asgi_request(
                "POST",
                "/api/documents",
                body=multipart_body,
                headers={**headers, "Content-Type": multipart_content_type},
            )
        )
        document = _json_response(upload_response)
        document_id = str(document.get("document_id") or "")

        if document_id:
            process_response = asyncio.run(
                _asgi_request("POST", f"/api/documents/{document_id}/process", headers=headers)
            )
            job = _json_response(process_response)
            quality_response = asyncio.run(
                _asgi_request("GET", f"/api/documents/{document_id}/quality", headers=headers)
            )
            quality = _json_response(quality_response)
            tenant_settings = settings_for_tenant(settings, tenant_id)
            for export_format in export_formats:
                export_response = asyncio.run(
                    _asgi_request(
                        "GET",
                        f"/api/documents/{document_id}/export?format={export_format}",
                        headers=headers,
                    )
                )
                extension = EXPORT_EXTENSIONS.get(export_format, export_format)
                persisted = tenant_settings.exports_dir / f"{document_id}.{extension}"
                exports.append(
                    {
                        "format": export_format,
                        "status_code": export_response.status_code,
                        "exists": export_response.status_code == 200 and persisted.is_file(),
                    }
                )
        else:
            job = {}
            quality = {}

    audit_summary = _summarize_audit(settings, tenant_id, actor=actor, document_id=document_id, since=run_started_at)
    required_audit_actions = {"auth.denied", "document.upload", "document.process", "document.export"}
    auth_denial_passed = denied_response.status_code in {401, 403} and "auth.denied" in audit_summary["actions"]
    tenant_header_required_passed = missing_tenant_response.status_code == 400
    audit_passed = required_audit_actions.issubset(set(audit_summary["actions"]))
    failed_exports = [item["format"] for item in exports if not item["exists"]]
    passed = (
        upload_response.status_code == 200
        and job.get("status") == "completed"
        and bool(quality.get("passed"))
        and not failed_exports
        and auth_denial_passed
        and tenant_header_required_passed
        and audit_passed
    )
    return {
        "report_type": "private_release_smoke",
        "generated_at": utc_now_iso(),
        "repo_commit": _git_commit(PROJECT_ROOT),
        "passed": passed,
        "sample_filename": sample_path.name,
        "data_dir_name": data_dir.name,
        "tenant_id": tenant_id,
        "document": {
            "document_id": document.get("document_id"),
            "file_type": document.get("file_type"),
            "source_system": document.get("source_system"),
            "profile_id": document.get("profile_id"),
        },
        "job": {
            "job_id": job.get("job_id"),
            "status": job.get("status"),
            "progress": job.get("progress"),
            "message": job.get("message"),
        },
        "quality": {
            "passed": quality.get("passed"),
            "score": quality.get("score"),
            "chunk_count": quality.get("chunk_count"),
            "issue_count": quality.get("issue_count"),
        },
        "required_export_formats": list(export_formats),
        "exports": exports,
        "failed_exports": failed_exports,
        "http": {
            "unauthorized_upload_status_code": denied_response.status_code,
            "missing_tenant_upload_status_code": missing_tenant_response.status_code,
            "authorized_upload_status_code": upload_response.status_code,
        },
        "audit": {
            "passed": audit_passed,
            "auth_denial_passed": auth_denial_passed,
            "tenant_header_required_passed": tenant_header_required_passed,
            "required_actions": sorted(required_audit_actions),
            "actions": audit_summary["actions"],
            "record_count": audit_summary["record_count"],
            "files_checked": audit_summary["files_checked"],
        },
    }


def write_synthetic_smoke_docx(path: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate a synthetic smoke sample.") from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("합성 비공개 릴리스 스모크 규정")
    doc.add_paragraph("제1조(목적) 이 규정은 비공개 릴리스 스모크 테스트를 위한 합성 문서이다.")
    doc.add_paragraph("제2조(처리) 문서 업로드, 전처리, 품질 확인, export를 검증한다.")
    doc.add_paragraph("제3조(MCP) 승인된 데이터는 로컬 규정 DB와 MCP 서버 연결 검증에 사용할 수 있다.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "검증"
    table.cell(1, 0).text = "MCP"
    table.cell(1, 1).text = "승인 후 연결"
    doc.save(path)
    return path


def _multipart_upload_body(sample_path: Path, fields: dict[str, str]) -> tuple[bytes, str]:
    boundary = f"reg-rag-smoke-{uuid4().hex}"
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    chunks.extend(
        [
            f"--{boundary}\r\n".encode("utf-8"),
            f'Content-Disposition: form-data; name="file"; filename="{sample_path.name}"\r\n'.encode("utf-8"),
            b"Content-Type: application/octet-stream\r\n\r\n",
            sample_path.read_bytes(),
            b"\r\n",
            f"--{boundary}--\r\n".encode("utf-8"),
        ]
    )
    return b"".join(chunks), f"multipart/form-data; boundary={boundary}"


async def _asgi_request(
    method: str,
    target: str,
    *,
    headers: dict[str, str] | None = None,
    body: bytes = b"",
) -> AsgiResponse:
    split = urlsplit(target)
    raw_headers = [
        (b"host", b"testserver"),
        (b"content-length", str(len(body)).encode("ascii")),
    ]
    for key, value in (headers or {}).items():
        raw_headers.append((key.lower().encode("latin-1"), str(value).encode("latin-1")))
    scope = {
        "type": "http",
        "asgi": {"version": "3.0", "spec_version": "2.3"},
        "http_version": "1.1",
        "method": method.upper(),
        "scheme": "http",
        "path": split.path,
        "raw_path": split.path.encode("ascii"),
        "query_string": split.query.encode("ascii"),
        "headers": raw_headers,
        "client": ("127.0.0.1", 12345),
        "server": ("testserver", 80),
    }
    request_sent = False
    status_code = 500
    response_headers: dict[str, str] = {}
    response_body: list[bytes] = []

    async def receive() -> dict[str, Any]:
        nonlocal request_sent
        if request_sent:
            return {"type": "http.disconnect"}
        request_sent = True
        return {"type": "http.request", "body": body, "more_body": False}

    async def send(message: dict[str, Any]) -> None:
        nonlocal status_code, response_headers
        if message["type"] == "http.response.start":
            status_code = int(message["status"])
            response_headers = {
                key.decode("latin-1").lower(): value.decode("latin-1")
                for key, value in message.get("headers", [])
            }
        elif message["type"] == "http.response.body":
            response_body.append(message.get("body", b""))

    await app_main.app(scope, receive, send)
    return AsgiResponse(status_code=status_code, headers=response_headers, content=b"".join(response_body))


def _json_response(response: AsgiResponse) -> dict[str, Any]:
    try:
        payload = response.json()
    except (json.JSONDecodeError, UnicodeDecodeError):
        return {}
    return payload if isinstance(payload, dict) else {}


def _summarize_audit(
    settings: Settings,
    tenant_id: str,
    *,
    actor: str,
    document_id: str,
    since: datetime,
) -> dict[str, Any]:
    sources = [
        ("base/repository/api_audit.jsonl", api_audit_path(settings)),
        ("tenant/repository/api_audit.jsonl", api_audit_path(settings_for_tenant(settings, tenant_id))),
    ]
    records: list[dict[str, Any]] = []
    files_checked: list[str] = []
    for label, path in sources:
        if not path.is_file():
            continue
        files_checked.append(label)
        for line in path.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            try:
                record = json.loads(line)
            except json.JSONDecodeError:
                continue
            if isinstance(record, dict) and _record_matches_current_smoke(
                record,
                tenant_id=tenant_id,
                actor=actor,
                document_id=document_id,
                since=since,
            ):
                records.append(record)
    actions = sorted({str(record.get("action", "")) for record in records if record.get("action")})
    return {
        "actions": actions,
        "record_count": len(records),
        "files_checked": files_checked,
    }


def _record_matches_current_smoke(
    record: dict[str, Any],
    *,
    tenant_id: str,
    actor: str,
    document_id: str,
    since: datetime,
) -> bool:
    created_at = _parse_datetime(record.get("created_at"))
    if created_at is None or created_at < since:
        return False
    action = str(record.get("action") or "")
    if action == "auth.denied":
        return (
            str(record.get("actor") or "") in {"unknown", actor}
            and str(record.get("claimed_tenant_id") or "") in {tenant_id, "default"}
        )
    if str(record.get("actor") or "") != actor:
        return False
    if str(record.get("tenant_id") or "") != tenant_id:
        return False
    if action in {"document.process", "document.export"} and document_id:
        return str(record.get("document_id") or "") == document_id
    return action == "document.upload"


def _parse_datetime(value: Any) -> datetime | None:
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except (TypeError, ValueError):
        return None
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _git_commit(root: Path) -> str:
    try:
        completed = subprocess.run(
            ["git", "-C", str(root), "rev-parse", "HEAD"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except FileNotFoundError:
        return ""
    if completed.returncode != 0:
        return ""
    return completed.stdout.decode("utf-8", "replace").strip()


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Run a local private-release sample smoke test.")
    sample_group = parser.add_mutually_exclusive_group()
    sample_group.add_argument(
        "--sample",
        type=Path,
        default=None,
        help=(
            "Approved sample document to upload and process. Relative paths resolve from the current working "
            "directory. Defaults to data/private_release_runtime/approved_samples/regulation_sample.hwp when "
            "--synthetic-sample is not used."
        ),
    )
    sample_group.add_argument(
        "--synthetic-sample",
        action="store_true",
        help="Generate a temporary synthetic DOCX sample so public/source-only clones can run the smoke.",
    )
    parser.add_argument("--data-dir", type=Path, default=None)
    parser.add_argument("--out-json", type=Path, default=None)
    parser.add_argument("--actor", default="private-release-smoke")
    parser.add_argument("--tenant-id", default="tenant-smoke")
    parser.add_argument(
        "--export-format",
        action="append",
        default=None,
        help="Export format to verify; may be repeated.",
    )
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.synthetic_sample:
        with tempfile.TemporaryDirectory(prefix="reg-rag-synthetic-sample-") as tmp:
            sample_path = write_synthetic_smoke_docx(Path(tmp) / "synthetic_private_release_smoke.docx")
            return _run_main_with_sample(args, sample_path, synthetic_sample=True)
    sample_arg = args.sample or DEFAULT_SAMPLE_PATH
    sample_path = (sample_arg if sample_arg.is_absolute() else Path.cwd() / sample_arg).resolve()
    return _run_main_with_sample(args, sample_path, synthetic_sample=False)


def _run_main_with_sample(args: argparse.Namespace, sample_path: Path, *, synthetic_sample: bool) -> int:
    if not sample_path.is_file():
        report = {
            "report_type": "private_release_smoke",
            "generated_at": utc_now_iso(),
            "repo_commit": _git_commit(PROJECT_ROOT),
            "passed": False,
            "synthetic_sample": synthetic_sample,
            "error": f"sample file not found: {sample_path}",
        }
        _emit_report(report, args.out_json)
        return 2

    export_formats = tuple(args.export_format or DEFAULT_EXPORT_FORMATS)
    if args.data_dir:
        report = run_smoke(
            sample_path=sample_path,
            data_dir=args.data_dir.resolve(),
            export_formats=export_formats,
            actor=args.actor,
            tenant_id=args.tenant_id,
        )
        report["data_dir_mode"] = "explicit"
        report["handoff_evidence"] = not synthetic_sample
    else:
        with tempfile.TemporaryDirectory(prefix="reg-rag-private-smoke-") as tmp:
            report = run_smoke(
                sample_path=sample_path,
                data_dir=Path(tmp) / "data",
                export_formats=export_formats,
                actor=args.actor,
                tenant_id=args.tenant_id,
            )
            report["data_dir_mode"] = "temporary"
            report["handoff_evidence"] = False
    report["synthetic_sample"] = synthetic_sample

    _emit_report(report, args.out_json)
    return 0 if report["passed"] else 1


def _emit_report(report: dict[str, Any], out_json: Path | None) -> None:
    output = json.dumps(report, ensure_ascii=False, indent=2)
    if out_json:
        out_json.parent.mkdir(parents=True, exist_ok=True)
        out_json.write_text(output + "\n", encoding="utf-8")
    print(output)


if __name__ == "__main__":
    raise SystemExit(main())
