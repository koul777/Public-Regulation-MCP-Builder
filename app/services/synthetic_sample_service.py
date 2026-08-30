from __future__ import annotations

from functools import lru_cache
from io import BytesIO


SYNTHETIC_SAMPLE_FILENAME = "synthetic_beginner_regulation.docx"
SYNTHETIC_SAMPLE_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@lru_cache(maxsize=1)
def build_synthetic_regulation_docx() -> bytes:
    """Build a redistributable, institution-free DOCX for first-run practice."""

    from docx import Document

    document = Document()
    document.core_properties.title = "합성 복무규정"
    document.core_properties.subject = "초보자 전처리 연습용 합성 문서"
    document.core_properties.author = "Public Regulation MCP Builder"
    document.add_heading("합성 복무규정", level=1)
    document.add_paragraph(
        "이 문서는 프로그램 사용법을 연습하기 위해 자동 생성한 합성 자료이며 실제 기관 규정이 아닙니다."
    )
    document.add_paragraph(
        "제1조(목적) 이 규정은 합성 예시 문서의 휴가 신청과 검토 절차를 정함을 목적으로 한다."
    )
    document.add_paragraph(
        "제2조(휴가 신청) ① 직원은 휴가 시작 전에 신청서를 제출한다. "
        "② 긴급한 사유가 있으면 담당자에게 먼저 알리고 신청서를 보완한다."
    )
    document.add_paragraph(
        "제3조(승인 및 기록) 담당자는 신청 내용을 확인한 뒤 승인 여부와 처리 결과를 기록한다."
    )
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "구분"
    table.cell(0, 1).text = "합성 처리 기준"
    table.cell(1, 0).text = "일반 휴가"
    table.cell(1, 1).text = "신청서 제출 후 담당자 확인"
    table.cell(2, 0).text = "긴급 휴가"
    table.cell(2, 1).text = "담당자에게 먼저 알린 뒤 신청서 보완"
    document.add_paragraph("부칙 이 합성 규정은 연습을 시작한 날부터 적용한다.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
